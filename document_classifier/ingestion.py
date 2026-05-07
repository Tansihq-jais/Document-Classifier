"""Ingestion module for the document-classifier pipeline.

Routes files to the correct extraction strategy and returns clean text.

Supported formats:
    - .pdf  → pdfplumber; falls back to pytesseract if text < pdf_text_min_chars chars
    - .png / .jpg / .jpeg → pytesseract with lang="eng+hin"
    - .docx → python-docx
    - other → IngestionError(error_type="unsupported_format")

All extraction exceptions are caught and returned as IngestionError(error_type="corrupt_file").
If pytesseract returns an empty string for a non-empty image, IngestionError(error_type="ocr_failure")
is returned.

For classification, use extract_for_classification() which calls the remote API
(staging-api.gipdataboard.io) for PDFs and images, falling back to local extraction
if the API is unavailable or disabled.

Raw document text is never written to logs.
"""

from __future__ import annotations

import os
import time
from dataclasses import dataclass, field

import pdfplumber  # type: ignore[import]
import pytesseract  # type: ignore[import]
import docx  # type: ignore[import]
from PIL import Image  # type: ignore[import]

# Point pytesseract at the Tesseract binary on Windows if it is not on PATH.
import os as _os
_TESSERACT_DEFAULT = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
if _os.name == "nt" and _os.path.isfile(_TESSERACT_DEFAULT):
    pytesseract.pytesseract.tesseract_cmd = _TESSERACT_DEFAULT

try:
    from pdf2image import convert_from_path  # type: ignore[import]
    _PDF2IMAGE_AVAILABLE = True
except ImportError:  # pragma: no cover
    _PDF2IMAGE_AVAILABLE = False
    convert_from_path = None  # type: ignore[assignment]

from document_classifier.config import PipelineConfig
from document_classifier.logger import log_stage_complete, log_stage_error

# ---------------------------------------------------------------------------
# Module-level configuration (can be overridden via configure_ingestion)
# ---------------------------------------------------------------------------

_config = PipelineConfig()

SUPPORTED_FORMATS = [".pdf", ".png", ".jpg", ".jpeg", ".docx"]


def configure_ingestion(config: PipelineConfig) -> None:
    """Override the module-level PipelineConfig.

    Args:
        config: A PipelineConfig instance to use for all subsequent extract() calls.
    """
    global _config
    _config = config


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------


@dataclass
class ExtractionResult:
    """Successful extraction result from any supported format.

    Attributes:
        text: Extracted text (may be mixed Hindi/English).
        source_file: Original file path.
        file_format: One of "pdf", "image", "docx".
        extraction_method: One of "pdfplumber", "ocr", "docx_parser".
        page_count: Number of pages (1 for images/docx when not applicable).
    """

    text: str
    source_file: str
    file_format: str
    extraction_method: str
    page_count: int


@dataclass
class IngestionError:
    """Structured error returned when extraction fails.

    Attributes:
        file_path: Path of the file that caused the error.
        error_type: One of "unsupported_format", "corrupt_file", "ocr_failure".
        message: Human-readable description of the failure.
        supported_formats: Populated only for error_type="unsupported_format".
    """

    file_path: str
    error_type: str
    message: str
    supported_formats: list[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def extract(file_path: str) -> ExtractionResult | IngestionError:
    """Extract text from a document file.

    Routes the file to the appropriate extraction strategy based on its extension.
    Logs stage completion or errors via logger.py.

    Args:
        file_path: Absolute or relative path to the document file.

    Returns:
        ExtractionResult on success, IngestionError on failure.
    """
    start_time = time.monotonic()
    # Use the filename as a pseudo document_id for logging (no UUID at this stage)
    document_id = os.path.basename(file_path)

    ext = os.path.splitext(file_path)[1].lower()

    if ext not in SUPPORTED_FORMATS:
        error = IngestionError(
            file_path=file_path,
            error_type="unsupported_format",
            message=(
                f"Unsupported file format '{ext}'. "
                f"Supported formats: {', '.join(SUPPORTED_FORMATS)}"
            ),
            supported_formats=list(SUPPORTED_FORMATS),
        )
        log_stage_error(
            stage="ingestion",
            document_id=document_id,
            error_type=error.error_type,
            message=error.message,
        )
        return error

    try:
        if ext == ".pdf":
            result = _extract_pdf(file_path)
        elif ext in (".png", ".jpg", ".jpeg"):
            result = _extract_image(file_path)
        else:  # .docx
            result = _extract_docx(file_path)
    except Exception as exc:  # noqa: BLE001
        error = IngestionError(
            file_path=file_path,
            error_type="corrupt_file",
            message=f"Failed to extract text from '{file_path}': {exc}",
        )
        log_stage_error(
            stage="ingestion",
            document_id=document_id,
            error_type=error.error_type,
            message=error.message,
        )
        return error

    if isinstance(result, IngestionError):
        log_stage_error(
            stage="ingestion",
            document_id=document_id,
            error_type=result.error_type,
            message=result.message,
        )
        return result

    duration_ms = (time.monotonic() - start_time) * 1000.0
    log_stage_complete(
        stage="ingestion",
        document_id=document_id,
        duration_ms=duration_ms,
        metadata={
            "file_format": result.file_format,
            "extraction_method": result.extraction_method,
            "page_count": result.page_count,
        },
    )
    return result


# ---------------------------------------------------------------------------
# Internal extraction helpers
# ---------------------------------------------------------------------------


def _call_extraction_api(
    file_path: str,
    ext: str,
) -> str | None:
    """Call the remote extraction API and return text, or None on failure.

    - PDFs  → multipart/form-data  (API expects a file field)
    - Images → raw binary body     (API expects application/octet-stream)
    - Other formats are not sent to the API.

    Args:
        file_path: Path to the file to extract text from.
        ext: Lowercase file extension including dot (e.g. ".pdf", ".jpg").

    Returns:
        Extracted text string on success, None if the API is disabled,
        unreachable, or returns an error.
    """
    api_url = _config.extraction_api_url
    if not api_url:
        return None

    try:
        import requests  # type: ignore[import]

        timeout = _config.extraction_api_timeout

        with open(file_path, "rb") as f:
            file_bytes = f.read()

        if ext == ".pdf":
            # PDFs: multipart/form-data with field name "file"
            response = requests.post(
                api_url,
                files={"file": (os.path.basename(file_path), file_bytes, "application/pdf")},
                timeout=timeout,
            )
        elif ext in (".jpg", ".jpeg", ".png", ".docx"):
            # Images and DOCX: raw binary body
            response = requests.post(
                api_url,
                data=file_bytes,
                headers={"Content-Type": "application/octet-stream"},
                timeout=timeout,
            )
        else:
            return None

        if response.status_code == 200:
            data = response.json()
            if data.get("success") and data.get("text", "").strip():
                return data["text"]

    except Exception:  # noqa: BLE001
        # Network error, timeout, JSON parse error — fall through to local
        pass

    return None


def extract_for_classification(file_path: str) -> "ExtractionResult | IngestionError":
    """Extract text for classification using the remote API when available.

    For PDFs and images, tries the remote extraction API first.
    Falls back to local extraction (pdfplumber / pytesseract) if:
      - The API URL is not configured
      - The API call fails or times out
      - The API returns empty text

    .docx files always use local extraction (python-docx).

    Args:
        file_path: Absolute or relative path to the document file.

    Returns:
        ExtractionResult on success, IngestionError on failure.
    """
    start_time = time.monotonic()
    document_id = os.path.basename(file_path)
    ext = os.path.splitext(file_path)[1].lower()

    if ext not in SUPPORTED_FORMATS:
        error = IngestionError(
            file_path=file_path,
            error_type="unsupported_format",
            message=(
                f"Unsupported file format '{ext}'. "
                f"Supported formats: {', '.join(SUPPORTED_FORMATS)}"
            ),
            supported_formats=list(SUPPORTED_FORMATS),
        )
        log_stage_error(
            stage="ingestion",
            document_id=document_id,
            error_type=error.error_type,
            message=error.message,
        )
        return error

    # --- Try remote API for PDFs, images and DOCX ---
    if ext in (".pdf", ".jpg", ".jpeg", ".png", ".docx"):
        api_text = _call_extraction_api(file_path, ext)
        if api_text:
            file_format = "pdf" if ext == ".pdf" else ("docx" if ext == ".docx" else "image")
            duration_ms = (time.monotonic() - start_time) * 1000.0
            result = ExtractionResult(
                text=api_text,
                source_file=file_path,
                file_format=file_format,
                extraction_method="api",
                page_count=1,
            )
            log_stage_complete(
                stage="ingestion",
                document_id=document_id,
                duration_ms=duration_ms,
                metadata={
                    "file_format": file_format,
                    "extraction_method": "api",
                    "page_count": 1,
                },
            )
            return result

    # --- Fallback to local extraction ---
    return extract(file_path)


# ---------------------------------------------------------------------------
# Internal extraction helpers
# ---------------------------------------------------------------------------


def _extract_pdf(file_path: str) -> ExtractionResult | IngestionError:
    """Extract text from a PDF using pdfplumber, falling back to OCR if needed."""
    with pdfplumber.open(file_path) as pdf:
        page_count = len(pdf.pages)
        parts: list[str] = []
        for page in pdf.pages:
            page_text = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
            parts.append(page_text)
        text = "\n".join(parts)

    if len(text.strip()) < _config.pdf_text_min_chars:
        # Fall back to OCR
        return _ocr_pdf(file_path, page_count)

    return ExtractionResult(
        text=text,
        source_file=file_path,
        file_format="pdf",
        extraction_method="pdfplumber",
        page_count=page_count,
    )


def _ocr_pdf(file_path: str, page_count: int) -> ExtractionResult | IngestionError:
    """Run pytesseract OCR on a PDF (converted to images via pdf2image or pdfplumber)."""
    images = []

    # Try pdf2image first (requires Poppler)
    if convert_from_path is not None:
        try:
            images = convert_from_path(file_path)
        except Exception:
            images = []

    # Fallback: use pdfplumber to render each page as a PIL image
    if not images:
        try:
            import pdfplumber as _pdfplumber
            with _pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    img = page.to_image(resolution=150).original
                    images.append(img)
        except Exception:
            images = []

    # Last resort: try opening directly as image (single-page image PDFs)
    if not images:
        try:
            images = [Image.open(file_path)]
        except Exception:
            pass

    if not images:
        return IngestionError(
            file_path=file_path,
            error_type="corrupt_file",
            message=f"Could not render pages from '{file_path}' for OCR",
        )

    parts: list[str] = []
    
    for img in images:
        # Try preprocessed image first (best for most cases)
        preprocessed = _preprocess_image(img)
        text = pytesseract.image_to_string(preprocessed, lang="eng+hin", config="--psm 3")
        
        # Fallback: if preprocessing produced empty text, try raw image
        if not text.strip():
            text = pytesseract.image_to_string(img, lang="eng+hin", config="--psm 3")
        
        parts.append(text)

    text = "\n".join(parts)

    if not text.strip():
        return IngestionError(
            file_path=file_path,
            error_type="ocr_failure",
            message=f"OCR produced empty text for '{file_path}'",
        )

    return ExtractionResult(
        text=text,
        source_file=file_path,
        file_format="pdf",
        extraction_method="ocr",
        page_count=page_count,
    )


def _deskew_image(img: Image.Image) -> Image.Image:
    """Detect and correct image rotation/skew using fast projection profile method.
    
    This helps OCR handle tilted/rotated document photos.
    Uses a coarse-to-fine search to minimize computation time.
    
    Args:
        img: Input PIL image (should be grayscale).
    
    Returns:
        Deskewed PIL image.
    """
    import numpy as np
    from scipy import ndimage
    
    # Convert to numpy array
    img_array = np.array(img)
    
    # Ensure grayscale
    if len(img_array.shape) == 3:
        img_array = np.mean(img_array, axis=2).astype(np.uint8)
    
    # Binarize (Otsu's method approximation)
    threshold = np.mean(img_array)
    binary = img_array < threshold
    
    # Coarse search: try angles from -15 to +15 degrees in 5-degree steps
    # Most ID cards are within ±15 degrees of horizontal
    coarse_angles = np.arange(-15, 16, 5)
    best_score = -1
    best_angle = 0
    
    for angle in coarse_angles:
        # Rotate and compute projection profile variance
        rotated = ndimage.rotate(binary, angle, reshape=False, order=0)
        # Sum along horizontal axis (projection profile)
        projection = np.sum(rotated, axis=1)
        # Higher variance = better alignment (text lines create peaks)
        score = np.var(projection)
        if score > best_score:
            best_score = score
            best_angle = angle
    
    # Fine search: refine around best coarse angle in 1-degree steps
    fine_angles = np.arange(best_angle - 4, best_angle + 5, 1)
    for angle in fine_angles:
        rotated = ndimage.rotate(binary, angle, reshape=False, order=0)
        projection = np.sum(rotated, axis=1)
        score = np.var(projection)
        if score > best_score:
            best_score = score
            best_angle = angle
    
    # Only rotate if angle is significant (> 2 degrees)
    if abs(best_angle) > 2:
        # Rotate original image
        rotated = ndimage.rotate(img_array, best_angle, reshape=True, cval=255)
        return Image.fromarray(rotated.astype(np.uint8))
    
    return img


def _preprocess_image(img: Image.Image) -> Image.Image:
    """Preprocess a PIL image to improve OCR accuracy on scanned ID cards.

    Steps:
    1. Convert to RGB to normalise mode (handles RGBA/palette images).
    2. Upscale if either dimension is below 1500 px — Tesseract performs best
       at ~300 DPI; scanned cards on A4 pages are often much smaller.
    3. Apply a modest contrast boost (1.5×) and sharpness boost (1.5×) to
       make text edges crisper without washing out light-coloured text.
    4. Convert to greyscale — reduces noise and speeds up Tesseract.

    Args:
        img: Input PIL image in any mode.

    Returns:
        Preprocessed greyscale PIL image ready for pytesseract.
    """
    from PIL import ImageEnhance  # already available via Pillow

    img = img.convert("RGB")

    # Upscale so the card region has enough resolution for Tesseract
    w, h = img.size
    if w < 1500 or h < 1500:
        scale = max(1500 / w, 1500 / h)
        new_size = (int(w * scale), int(h * scale))
        img = img.resize(new_size, Image.LANCZOS)

    # Modest enhancement — aggressive values (2.0+) wash out light-on-light text
    img = ImageEnhance.Contrast(img).enhance(1.5)
    img = ImageEnhance.Sharpness(img).enhance(1.5)

    # Greyscale reduces colour noise and is sufficient for text extraction
    img = img.convert("L")
    
    # Deskew to correct rotation/tilt (with error handling)
    try:
        img = _deskew_image(img)
    except Exception:
        # If deskewing fails, continue with non-deskewed image
        pass

    return img


def _extract_image(file_path: str) -> ExtractionResult | IngestionError:
    """Extract text from an image file using pytesseract OCR.
    
    Uses preprocessed image (upscaled + enhanced + deskewed) with PSM 3 (fully automatic
    page segmentation) which provides the best balance of accuracy and speed for ID cards.
    Falls back to raw image if preprocessing produces empty text.
    """
    img = Image.open(file_path)

    # Try preprocessed image first (best for most cases)
    preprocessed = _preprocess_image(img)
    text = pytesseract.image_to_string(preprocessed, lang="eng+hin", config="--psm 3")

    # Fallback: if preprocessing produced empty text, try raw image
    if not text.strip():
        text = pytesseract.image_to_string(img, lang="eng+hin", config="--psm 3")
        method = "ocr_raw_psm3"
    else:
        method = "ocr_preprocessed_psm3"

    if not text.strip():
        return IngestionError(
            file_path=file_path,
            error_type="ocr_failure",
            message=f"OCR produced empty text for '{file_path}'",
        )

    return ExtractionResult(
        text=text,
        source_file=file_path,
        file_format="image",
        extraction_method=method,
        page_count=1,
    )


def _extract_docx(file_path: str) -> ExtractionResult:
    """Extract text from a .docx file using python-docx."""
    doc = docx.Document(file_path)
    parts: list[str] = []
    for para in doc.paragraphs:
        parts.append(para.text)
    # Also extract text from tables (tabular data as whitespace-delimited text)
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text for cell in row.cells]
            parts.append("\t".join(row_texts))

    text = "\n".join(parts)

    return ExtractionResult(
        text=text,
        source_file=file_path,
        file_format="docx",
        extraction_method="docx_parser",
        page_count=1,
    )
